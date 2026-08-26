import sys, re
from docx import Document
from docx.shared import Pt

def add_runs(paragraph, text):
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for p in parts:
        if len(p) >= 4 and p.startswith('**') and p.endswith('**'):
            r = paragraph.add_run(p[2:-2]); r.bold = True
        elif p:
            paragraph.add_run(p)

def convert(md_path, docx_path):
    doc = Document()
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        if line.strip().startswith('```'):
            i += 1
            code = []
            while i < n and not lines[i].strip().startswith('```'):
                code.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code))
            run.font.name = 'Consolas'; run.font.size = Pt(9)
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3); i += 1; continue
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2); i += 1; continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1); i += 1; continue
        if line.strip().startswith('|'):
            tbl = []
            while i < n and lines[i].strip().startswith('|'):
                tbl.append(lines[i]); i += 1
            rows = [[c.strip() for c in tl.strip().strip('|').split('|')] for tl in tbl]
            data = [r for r in rows if not all(set(c) <= set('-: ') for c in r)]
            if data:
                t = doc.add_table(rows=len(data), cols=len(data[0]))
                t.style = 'Table Grid'
                for ri, row in enumerate(data):
                    for ci, cell in enumerate(row):
                        add_runs(t.cell(ri, ci).paragraphs[0], cell)
            continue
        if line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet'); add_runs(p, line[2:]); i += 1; continue
        if line.strip() == '':
            i += 1; continue
        if line.startswith('> '):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(12)
            add_runs(p, line[2:]); i += 1; continue
        p = doc.add_paragraph(); add_runs(p, line); i += 1
    doc.save(docx_path)
    print('saved', docx_path)

if __name__ == '__main__':
    convert(sys.argv[1], sys.argv[2])
